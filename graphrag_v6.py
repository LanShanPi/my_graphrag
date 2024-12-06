from llama_index.core.prompts import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType
import os
from llama_index.core.indices.property_graph import LLMSynonymRetriever
from llama_index.core.indices.property_graph import TextToCypherRetriever
from llama_index.core import SimpleDirectoryReader, KnowledgeGraphIndex, StorageContext, load_index_from_storage
from llama_index.core import PropertyGraphIndex
from llama_index.core.graph_stores import SimpleGraphStore
from llama_index.llms.openai import OpenAI
from pyvis.network import Network
from llama_index.llms.openai import OpenAI
import openai
import torch
from llama_index.core.schema import Document
import re
from docx import Document as DocxDocument
import fitz  # PyMuPDF
from config import OPENAI_API_KEY1 as OPENAI_API_KEY
from config import API_BASE1 as API_BASE
from config import API_BASE2
from prompt.prompt import hongloumeng as triplet_extraction_template
from functools import lru_cache
import networkx as nx
from matplotlib import pyplot as plt
from llama_index.core.schema import Node
from llama_index.core.schema import MediaResource
from openai import OpenAI as local_openai


# 设置 OpenAI API
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
openai.api_key = os.environ["OPENAI_API_KEY"]
openai.api_base = API_BASE

# 配置 GPU/CPU 设备
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
llm = OpenAI(temperature=0, model="gpt-4o", device=device)

# 定义统一的存储目录
def setup_storage_dir(person_name):
    """
    创建统一存储目录
    :param person_name: 目标人物名称
    :return: 存储目录路径
    """
    base_dir = os.path.join(os.getcwd(), f"{person_name}_knowledge_graph")
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
    return base_dir


# 文件读取函数
def read_docx(file_path):
    docx = DocxDocument(file_path)
    content = []
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
    return content

def preprocess_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def chunk_text_with_overlap(text, chunk_size, overlap):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def read_pdf_with_pymupdf(file_path):
    documents = []
    pdf = fitz.open(file_path)
    for i, page in enumerate(pdf):
        text = page.get_text("text")
        if text:
            processed_text = preprocess_text(text)
            documents.append(Document(text=processed_text, extra_info={"source": f"Page {i + 1}"}))
    pdf.close()
    return documents

def process_file(file_path, file_type):
    if file_type == "docx":
        docx_content = read_docx(file_path)
        if not docx_content:
            raise ValueError("未从文件中读取到任何内容，请检查文件格式或内容是否为空。")
        documents = [
            Document(text=preprocess_text(text), extra_info={"source": f"Content {i + 1}"})
            for i, text in enumerate(docx_content)
        ]
    elif file_type == "pdf":
        documents = read_pdf_with_pymupdf(file_path)
    elif file_type == "txt":
        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return documents

# 定义知识提取模板
def get_config():
    CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = PromptTemplate(
        triplet_extraction_template["CUSTOM_KG_TRIPLET_EXTRACT_TMPL"],
        prompt_type=PromptType.KNOWLEDGE_TRIPLET_EXTRACT
    )
    allowed_entity_types = triplet_extraction_template["allowed_entity_types"]
    allowed_relation_types = triplet_extraction_template["allowed_relation_types"]
    return allowed_entity_types, allowed_relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT

# 过滤特定人物的三元组
def filter_triplets_by_person(triplets, person_name):
    return [
        triplet for triplet in triplets
        if person_name in triplet["head"] or person_name in triplet["tail"]
    ]

def check_subfolder_exists(parent_folder, subfolder_name):
    """
    检查指定文件夹中是否存在某个子文件夹
    :param parent_folder: 父文件夹路径
    :param subfolder_name: 子文件夹名称
    :return: True 如果存在，False 如果不存在
    """
    subfolder_path = os.path.join(parent_folder, subfolder_name)
    return os.path.exists(subfolder_path) and os.path.isdir(subfolder_path)



def generate_knowledge_graph(file_path,file_type,graph_name,storage_dir):
    # 文件处理和知识图谱构建
    documents = process_file(file_path, file_type)
    chunk_size = 512
    overlap = 50
    chunked_documents = []
    for doc in documents:
        text_chunks = chunk_text_with_overlap(doc.text, chunk_size, overlap)
        for chunk in text_chunks:
            chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))

    # 加载模板
    entity_types, relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = get_config()
    # 创建存储上下文
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    # 构建针对特定人物的知识图谱
    index = PropertyGraphIndex.from_documents(
        chunked_documents,
        max_triplets_per_chunk=10,
        storage_context=storage_context,
        show_progress=True,
        include_embeddings=False,
        kg_triple_extract_template=CUSTOM_KG_TRIPLET_EXTRACT_PROMPT,
        allowed_entity_types=entity_types,
        allowed_relation_types=relation_types,
        triplet_filter_fn=lambda triplets: filter_triplets_by_person(triplets, graph_name)
    )

    # 存储知识图谱和可视化文件到统一目录
    index_dir = os.path.join(storage_dir, "index")
    if not os.path.exists(index_dir):
        os.makedirs(index_dir)
    storage_context.persist(persist_dir=index_dir)
    # 生成html文件
    html_file = os.path.join(storage_dir, f"{graph_name}_graph.html")
    index.property_graph_store.save_networkx_graph(name=html_file)
    print("html文件已生成")
    
    return index

def load_knowledge_graph(storage_dir):
    # 加载已存储的索引
    print("已存在相关知识库")
    storage_context = StorageContext.from_defaults(persist_dir=storage_dir+"/"+"index")
    index = load_index_from_storage(storage_context)
    return index



def generate_subgraph(index, store_dir, noun=None):
    """
    从原始知识图谱中抽取子图并存储。
    :param index: 知识图谱索引
    :param store_dir: 存储路径
    :param noun: 查询的名词
    """
    # 生成noun别的称呼
    client = local_openai(  
    api_key=OPENAI_API_KEY,
    base_url=API_BASE2 
)
    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "根据你对中国历史古籍的理解，回答问题"},
            {"role": "user","content": f"{noun}有哪些别的称呼，请尽可能全面的罗列，显示格式为：&名字1&名字2...名字n&"}
        ]
    )
    nouns = completion.choices[0].message.content.split("&")[1:-1]
    nouns.append(noun)

    # 从知识图谱中直接检索与名词相关的三元组
    triplets = list(index.property_graph_store.get_triplets(entity_names=nouns))
    # 如果没有找到三元组，输出提示信息
    if not triplets:
        print(f"未找到与名词 '{noun}' 相关的三元组。请检查名词是否存在于知识图谱中。")
        return
    
    # 创建NetworkX图
    G = nx.DiGraph()
    # 添加节点和边
    for subj, rel, obj in triplets:
        G.add_node(subj.id, label=subj.name or f"节点-{subj.id}")
        G.add_node(obj.id, label=obj.name or f"节点-{obj.id}")
        G.add_edge(subj.id, obj.id, label=rel.label or f"关系-{rel.id}")
    # 使用pyvis可视化
    net = Network(notebook=True, cdn_resources="in_line", directed=True)
    net.from_nx(G)
    # # 生成静态图
    # net.save_graph(os.path.join(store_dir, "static_subgraph.html"))
    # 生成动态图
    net.show(os.path.join(store_dir, "subgraph.html"))
    print(f"子图已存储到 {os.path.join(store_dir, 'subgraph.html')}")

    # # 提取节点ID和文本
    # nodes = []
    # for subj, rel, obj in triplets:
    #     # 假设subj和obj是EntityNode对象，并且包含id和text属性
    #     nodes.append(Node(id=subj.id, text=MediaResource(content=subj.name or f"节点-{subj.id}")))
    #     nodes.append(Node(id=rel.id, text=MediaResource(content=rel.label or f"关系-{rel.id}")))
    #     nodes.append(Node(id=obj.id, text=MediaResource(content=obj.name or f"节点-{obj.id}")))

    
    # # 确保节点列表不为空
    # if not nodes:
    #     raise ValueError("未找到与该名词相关的节点。")

    # # 创建存储上下文
    # storage_context = StorageContext.from_defaults()

    # # 初始化 PropertyGraphIndex
    # subgraph_index = PropertyGraphIndex(nodes=nodes, storage_context=storage_context)

    # # 保存子图索引
    # subgraph_dir = os.path.join(store_dir, "subgraph", "index")
    # subgraph_index.storage_context.persist(persist_dir=subgraph_dir)

    # print(f"子图已存储到 {subgraph_dir}")



def get_response(index,queries):
    # queries为问题列表
    query_engine = index.as_query_engine(llm=llm, include_text=False)
    # query_engine.set_query_mode("compact")
    response = query_engine.query(queries)
    return response


if __name__ == "__main__":
    file_path = "/home/share/shucshqyfzyxgsi/home/lishuguang/my_graphrag/data/红.txt"
    file_type = "txt"
    graph_name = "hongloumeng"
    
    # 生成存储路径
    storage_dir = setup_storage_dir(graph_name)
    # 若未生成相关图谱，则进行生成，若之前已生成则直接加载
    if not check_subfolder_exists(storage_dir, "index"):
        index = generate_knowledge_graph(file_path,file_type,graph_name,storage_dir)
    else:
        index = load_knowledge_graph(storage_dir)

    
    # response = get_response(index,queries=["朱元璋的有哪些别名"])
    response = get_response(index,queries="王熙凤是谁")
    print(response)
    # generate_subgraph(index,storage_dir,noun="朱元璋")
