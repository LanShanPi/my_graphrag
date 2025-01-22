from llama_index.core.prompts import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType
from llama_index.embeddings.openai import OpenAIEmbedding
import os
from llama_index.core import SimpleDirectoryReader, StorageContext, load_index_from_storage
from llama_index.core import PropertyGraphIndex
from llama_index.core.graph_stores import SimpleGraphStore
from llama_index.llms.openai import OpenAI as LlamaOpenAI
from pyvis.network import Network
import openai
import torch 
from llama_index.core.schema import Document
import re
from docx import Document as DocxDocument
import fitz  # PyMuPDF
from config import OPENAI_API_KEY2 as OPENAI_API_KEY
from config import API_BASE1 as API_BASE
from prompt.prompt import novel_kg_extraction as triplet_extraction_template
from functools import lru_cache
import networkx as nx
from matplotlib import pyplot as plt
from openai import OpenAI as local_openai
from prompt.response_prompt import mingchao_person as prompt_
import os
import re
import time
# from llama_index.core.extractors import BaseExtractor
from llama_index.core.extractors.metadata_extractors import BaseExtractor
from llama_index.core.ingestion import IngestionPipeline
# from llama_index.core.vector_stores.types import MetadataFilters, ExactMatchFilter
import logging



# # 设置 OpenAI API
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
openai.api_key = os.environ["OPENAI_API_KEY"]
openai.api_base = API_BASE

# 配置 GPU/CPU 设备
# device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
# llm = LlamaOpenAI(temperature=0, model="gpt-4o", device=device)
llm = LlamaOpenAI(temperature=0, model="gpt-4o", api_key=OPENAI_API_KEY, base_url=API_BASE,timeout=600)

# 定义统一的存储目录
def setup_storage_dir(file_name):
    """
    创建统一存储目录
    :param file_name: 文件名称
    :return: 存储目录路径
    """
    base_dir = os.path.join(os.getcwd(), f"{file_name}_knowledge_graph")
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
    return base_dir


# 文件读取函数
def read_docx(file_path):
    docx = DocxDocument(file_path)
    content = []
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
    return content

def preprocess_text(text):
    # 将字符串 text 中的所有连续空白字符替换为单个空格。
    return re.sub(r'\s+', ' ', text).strip()

def chunk_text_with_overlap(text, chunk_size, overlap):
    # 将字符串切割成指定大小的块
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def read_pdf_with_pymupdf(file_path):
    documents = []
    pdf = fitz.open(file_path)
    for i, page in enumerate(pdf):
        text = page.get_text("text")
        if text:
            processed_text = preprocess_text(text)
            documents.append(Document(text=processed_text, extra_info={"source": f"Page {i + 1}"}))
    pdf.close()
    return documents

def process_file(file_path, file_type):
    if file_type == "docx":
        docx_content = read_docx(file_path)
        if not docx_content:
            raise ValueError("未从文件中读取到任何内容，请检查文件格式或内容是否为空。")
        documents = [
            Document(text=preprocess_text(text), extra_info={"source": f"Content {i + 1}"})
            for i, text in enumerate(docx_content)
        ]
    elif file_type == "pdf":
        documents = read_pdf_with_pymupdf(file_path)
    elif file_type == "txt":
        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return documents

# 定义知识提取模板
def get_config():
    CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = PromptTemplate(
        triplet_extraction_template["CUSTOM_KG_TRIPLET_EXTRACT_TMPL"],
        prompt_type=PromptType.KNOWLEDGE_TRIPLET_EXTRACT
    )
    allowed_entity_types = triplet_extraction_template["allowed_entity_types"]
    allowed_relation_types = triplet_extraction_template["allowed_relation_types"]
    return allowed_entity_types, allowed_relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT

def check_subfolder_exists(parent_folder, subfolder_name):
    """
    检查指定文件夹中是否存在某个子文件夹
    :param parent_folder: 父文件夹路径
    :param subfolder_name: 子文件夹名称
    :return: True 如果存在，False 如果不存在
    """
    subfolder_path = os.path.join(parent_folder, subfolder_name)
    return os.path.exists(subfolder_path) and os.path.isdir(subfolder_path)

def generate_knowledge_graph(file_path, file_type, dir_name, storage_dir):
    # File processing and knowledge graph construction
    documents = process_file(file_path, file_type)
    chunk_size = 256
    overlap = 10
    chunked_documents = []
    for doc in documents:
        text_chunks = chunk_text_with_overlap(doc.text, chunk_size, overlap)
        for chunk in text_chunks:
            chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))

    # Load templates
    entity_types, relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = get_config()

    # Create storage context
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    
    embed_model = OpenAIEmbedding(model_name="text-embedding-ada-002",api_key=OPENAI_API_KEY, base_url=API_BASE)
    # embed_model=embed_model,  # 指定编码模型
    # include_embeddings=True,

    # Build knowledge graph
    index = PropertyGraphIndex.from_documents(
        chunked_documents,
        llm=llm,
        embed_model=embed_model,  # 指定编码模型
        include_embeddings=True,
        max_triplets_per_chunk=10,
        storage_context=storage_context,
        show_progress=True,
        kg_triple_extract_template=CUSTOM_KG_TRIPLET_EXTRACT_PROMPT,
        allowed_entity_types=entity_types,
        allowed_relation_types=relation_types,
    )

    # Store knowledge graph and visualization files
    index_dir = os.path.join(storage_dir, "index")
    if not os.path.exists(index_dir):
        os.makedirs(index_dir)
    storage_context.persist(persist_dir=index_dir)

    # Convert knowledge graph to NetworkX graph
    G = nx.DiGraph()
    for triplet in index.property_graph_store.get_triplets():
        subj = triplet["head"]
        obj = triplet["tail"]
        rel = triplet["relation"]
        time = triplet.get("time", "未知时间")  # Get time or set default
        
        # Add nodes and edges to the graph
        G.add_node(subj, label=subj)
        G.add_node(obj, label=obj)
        # Add time to the edge label
        G.add_edge(subj, obj, label=f"{rel} ({time})")

    # Visualize with pyvis
    net = Network(notebook=True, cdn_resources="in_line", directed=True)
    net.from_nx(G)

    # Save HTML file
    html_file = os.path.join(storage_dir, f"{dir_name}_graph.html")
    net.show(html_file)
    print(f"HTML file has been generated: {html_file}")
    return index

def load_knowledge_graph(storage_dir):
    # 加载已存储的索引
    print("已存在相关知识库")
    storage_context = StorageContext.from_defaults(persist_dir=storage_dir+"/"+"index")
    index = load_index_from_storage(storage_context)
    return index


def get_response_v1(index,queries):
    ### 普通回复
    # 查询索引
    # include_text 用途：决定在查询结果中是否包含原始文本。
    # True：当设置为 True 时，查询结果将包含与匹配的三元组相关的完整文本块。这对于需要查看上下文或详细信息的应用场景非常有用。
    # False：当设置为 False 时，查询结果只包含匹配的三元组，而不包括完整的文本块。这可以减少返回的数据量，提高查询速度。
    ######
    # response_mode 用途：指定如何生成和组织查询响应。
    # tree_summarize：在这种模式下，系统会递归地构建一个树状结构来总结查询结果。这种模式适用于需要对大量信息进行总结的场景。
    # 工作原理：tree_summarize 会将所有相关的文本块合并，并通过多次调用 LLM 来生成一个总结，最终返回一个简洁的响应。
    # 适用场景：适合需要高层次总结的场景，结果更全面，但是速度可能慢

    # refine 模式用途：用于逐步改进和细化响应。
    # 工作原理：在 refine 模式下，系统会首先使用第一个文本块和查询生成一个初始答案。然后，它会将这个答案与下一个文本块结合，再次生成一个改进的答案。这个过程会持续进行，直到所有文本块都被处理完毕。
    # 适用场景：适合需要详细和精确答案的场景，因为它会逐步整合所有相关信息。

    # compact 模式用途：用于在生成响应之前合并文本块，以减少对 LLM 的调用次数。
    # 工作原理：在 compact 模式下，系统会尽可能多地将文本块合并到一个上下文窗口中，然后生成响应。这种方法减少了对 LLM 的调用次数，因为它在合并后的文本上进行处理。
    # 适用场景：适合需要快速响应的场景，因为它通过减少 LLM 调用次数来提高效率。

    # 这两种模式各有优劣，refine 提供更详细的答案，而 compact 提供更快的响应速度。

    # vector_store_query_mode：指定向量存储查询的模式，通常用于选择不同的向量检索策略。
        # DEFAULT：默认的向量搜索模式，通常用于基于向量相似度的检索。
        # SPARSE：稀疏向量搜索模式，适用于稀疏数据的检索。
        # HYBRID：混合搜索模式，结合了向量相似度和文本搜索。
        # TEXT_SEARCH：全文本搜索模式，基于文本内容进行检索。
        # SEMANTIC_HYBRID：语义混合模式，结合语义理解和向量相似度。

    # filters：用于在查询时应用元数据过滤器。可以根据文档的元数据（如作者、日期等）来筛选符合条件的节点。

    # # 定义元数据过滤器（只检索作者为“历史学家”在2023年的记录，key的值要根据原数据中的标记来设置，查看原数据中的标签暂未写）####
    # filters = MetadataFilters(
    #                 filters=[
    #                     ExactMatchFilter(key="author", value="历史学家"),
    #                     ExactMatchFilter(key="year", value=2023)
    #                     ]
    #                     )

    query_engine = index.as_query_engine(llm=llm, include_text=True,response_mode="tree_summarize",similarity_top_k=5)
    response = query_engine.query(queries)
    for idx, source in enumerate(response.source_nodes):
        # 使用 get_content() 方法获取节点内容
        print("[Source] " + str(idx) + ": ", source.node.get_content())

    return response


if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler("knowledge_graph.log"),
            logging.StreamHandler()
        ]
    )


    file_path = "/Users/hzl/Project/my_graphrag/data/高中物理测试题3.pdf"
    file_type = "pdf"
    dir_name = "highschool_physics"
    
    # 生成存储路径
    storage_dir = setup_storage_dir(dir_name)
    # 若未生成相关图谱，则进行生成，若之前已生成则直接加载
    if not check_subfolder_exists(storage_dir, "index"):
        index = generate_knowledge_graph(file_path,file_type,dir_name,storage_dir)
    else:
        index = load_knowledge_graph(storage_dir)

    input_ = None
    pre_prompt = "根据自身能力和检索到的知识尽可能详细的回复下述问题，且回复要满足：回答的准确性、回答的完整性、回答的逻辑性、回答的语言表达清晰性，这四个要求，仅需输出回答就好了，不需要额外的输出。下面是问题："
    while input_ != "over":
        input_ = input("请输入你针对’高中物理测试题3的‘这本书的问题：")
        start = time.time()
        print("#######################")
        response = get_response_v1(index,pre_prompt+input_)
        print("#######################")
        print("生成时间：",time.time()-start)
        print(response)
