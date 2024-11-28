
import os
from llama_index.core import SimpleDirectoryReader, KnowledgeGraphIndex, StorageContext, load_index_from_storage
from llama_index.core.graph_stores import SimpleGraphStore
from llama_index.llms.openai import OpenAI
from pyvis.network import Network
import IPython
import openai
import torch
from llama_index.readers.file import PyMuPDFReader
import pdfplumber
from llama_index.core.schema import Document
import re
from docx import Document as DocxDocument
from config import OPENAI_API_KEY1,API_BASE

# 设置 OpenAI API
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY1
openai.api_key = os.environ["OPENAI_API_KEY"]
openai.api_base = API_BASE


# 配置 OpenAI 模型，使用 GPU 加速
# temperature: 控制生成文本的随机性，值越低，生成的文本越确定
# model: 指定使用的 OpenAI 模型
# device: 指定模型运行的设备（CPU 或 GPU）
# torch.cuda.is_available()目前返回false，所以用的是cpu
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")  # 检查是否有可用的 GPU
llm = OpenAI(temperature=0, model="gpt-4o", device=device)  # 初始化 OpenAI 模型

# 读取 .docx 文件内容，包括段落和表格
def read_docx(file_path):
    """
    读取 .docx 文件，包括普通段落和表格内容
    :param file_path: .docx 文件路径
    :return: 提取的内容列表
    """
    docx = DocxDocument(file_path)
    content = []

    # 读取普通段落
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)

    # 读取表格内容
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))  # 将一行的单元格合并成一条记录

    return content

# 文本预处理和分块函数
def preprocess_text(text):
    """
    对文本进行清理，比如去除多余空格或特殊字符
    :param text: 原始文本
    :return: 清理后的文本
    """
    text = re.sub(r'\s+', ' ', text).strip()  # 清除多余空格
    return text

def chunk_text_with_overlap(text, chunk_size, overlap):
    """
    分块文本，支持重叠
    :param text: 原始文本
    :param chunk_size: 每块字符数
    :param overlap: 重叠字符数
    :return: 分块后的文本列表
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def custom_extract_triplets(text):
    """
    自定义三元组提取逻辑
    """
    triplets = []

    # 技术措施、管理措施提取
    pattern1 = re.findall(r'(\w+)\s+(是|适用于|用于|可以|能|提升|防治|增加|减少)\s+(\w+)', text)
    for match in pattern1:
        triplets.append((match[0], match[1], match[2]))

    # 结合表格行内容提取实体关系
    if " | " in text:  # 假设表格行用 " | " 分隔
        cells = text.split(" | ")
        if len(cells) >= 2:
            triplets.append((cells[0], "相关", cells[1]))

    return triplets


# 定义存储目录
# PERSIST_DIR = "honglou_knowledge_graph"
# HTML_FILE = "honglou_knowledge_graph.html"
PERSIST_DIR = "pingguo_knowledge_graph"
HTML_FILE = "pingguo_knowledge_graph.html"

# 检查存储目录是否存在
if not os.path.exists(PERSIST_DIR):
    file_path = "/home/share/shucshqyfzyxgsi/home/lishuguang/my_graphrag/红富士苹果生产.docx"
    # 读取 .docx 文件内容
    docx_content = read_docx(file_path)
    if not docx_content:
        raise ValueError("未从文件中读取到任何内容，请检查文件格式或内容是否为空。")

    # 转换内容为 Document 对象
    documents = []
    for i, text in enumerate(docx_content):
        processed_text = preprocess_text(text)
        documents.append(Document(text=processed_text, extra_info={"source": f"Content {i + 1}"}))


    ##########################################
    # # 加载txt文件
    # documents = SimpleDirectoryReader(input_files=["红楼梦.txt"]).load_data()


    # 分块文本
    chunk_size = 512  # 设置每个块的字符数
    overlap = 50  # 设置重叠长度
    chunked_documents = []
    for doc in documents:
        text_chunks = chunk_text_with_overlap(doc.text, chunk_size,overlap)
        for chunk in text_chunks:
            chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))

    # 创建存储上下文
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    
    # 创建 KnowledgeGraphIndex
    # max_triplets_per_chunk表示文档分块后每个块可以提取多少三元组（关系和事实）的数量
    # kg_triplet_extract_fn 允许用户根据特定的需求和数据特性调整三元组提取的方式，而不是依赖于默认的提取方法。
    # include_embeddings=True：语义搜索：通过包含嵌入，索引可以支持基于语义的查询，而不仅仅是关键词匹配。这使得查询能够返回语义上相关的结果，即使它们不包含查询中的确切词语。
                        # 相似性计算：嵌入允许计算不同文本之间的相似性，从而支持更复杂的查询和分析。
                        # 提高查询准确性：在某些应用场景中，嵌入可以提高查询的准确性和相关性，因为它们捕捉了文本的语义信息。
    # show_progress=True显示进度条
    index = KnowledgeGraphIndex.from_documents(
        chunked_documents,
        max_triplets_per_chunk=10,
        storage_context=storage_context,
        show_progress=True,
        include_embeddings=True,
        # kg_triplet_extract_fn=custom_extract_triplets,
    )
    # 存储索引
    storage_context.persist(persist_dir=PERSIST_DIR)
else:
    # 加载已存储的索引
    storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
    index = load_index_from_storage(storage_context)
# 查询索引
# include_text 用途：决定在查询结果中是否包含原始文本。
# True：当设置为 True 时，查询结果将包含与匹配的三元组相关的完整文本块。这对于需要查看上下文或详细信息的应用场景非常有用。
# False：当设置为 False 时，查询结果只包含匹配的三元组，而不包括完整的文本块。这可以减少返回的数据量，提高查询速度。
######
# response_mode 用途：指定如何生成和组织查询响应。
# tree_summarize：在这种模式下，系统会递归地构建一个树状结构来总结查询结果。这种模式适用于需要对大量信息进行总结的场景。
# 工作原理：tree_summarize 会将所有相关的文本块合并，并通过多次调用 LLM 来生成一个总结，最终返回一个简洁的响应。
# 适用场景：适合需要高层次总结的场景，结果更全面，但是速度可能慢

# refine 模式用途：用于逐步改进和细化响应。
# 工作原理：在 refine 模式下，系统会首先使用第一个文本块和查询生成一个初始答案。然后，它会将这个答案与下一个文本块结合，再次生成一个改进的答案。这个过程会持续进行，直到所有文本块都被处理完毕。
# 适用场景：适合需要详细和精确答案的场景，因为它会逐步整合所有相关信息。

# compact 模式用途：用于在生成响应之前合并文本块，以减少对 LLM 的调用次数。
# 工作原理：在 compact 模式下，系统会尽可能多地将文本块合并到一个上下文窗口中，然后生成响应。这种方法减少了对 LLM 的调用次数，因为它在合并后的文本上进行处理。
# 适用场景：适合需要快速响应的场景，因为它通过减少 LLM 调用次数来提高效率。

# 这两种模式各有优劣，refine 提供更详细的答案，而 compact 提供更快的响应速度。

query_engine = index.as_query_engine(llm=llm, include_text=False, response_mode="compact")
response = query_engine.query("红富士苹果生产面临哪些问题")
print(response)

# # 打印查询到的内容
# for node in response.source_nodes:
#     print(f"Node Content: {node.text}")

# 检查HTML文件是否存在
if not os.path.exists(HTML_FILE):
    # 获取 NetworkX 图
    g = index.get_networkx_graph()
    # 使用 Pyvis 可视化
    
    # notebook=True：
    # 用途：指定在 Jupyter Notebook 环境中显示图形。
    # 效果：当设置为 True 时，图形会直接嵌入在 Jupyter Notebook 的输出单元中，便于交互和查看。
    
    # cdn_resources="in_line"：
    # 用途：指定如何加载图形所需的 JavaScript 和 CSS 资源。
    # 效果："in_line" 表示将所有必要的资源内联到生成的 HTML 中，确保图形在没有外部网络连接的情况下也能正确显示。
    
    # directed=True：
    # 用途：指定生成的图形是否为有向图。
    # 效果：当设置为 True 时，图形中的边将带有方向，通常用箭头表示，适合表示有方向性的关系或流程。

    net = Network(
    notebook=True,  # 在 Jupyter Notebook 中显示
    cdn_resources="in_line",  # 内联资源
    directed=True,  # 有向图
    )
    net.from_nx(g)
    # 生成 HTML 文件
    net.show(HTML_FILE)
# 显示 HTML 文件
IPython.display.HTML(filename=HTML_FILE)