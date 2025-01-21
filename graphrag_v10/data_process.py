# data_processing.py

import os
import re
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from llama_index.core import SimpleDirectoryReader
from llama_index.core.schema import Document

def preprocess_text(text):
    """
    将字符串 text 中的所有连续空白字符替换为单个空格，并 strip。
    """
    return re.sub(r'\s+', ' ', text).strip()

def chunk_text_with_overlap(text, chunk_size, overlap):
    """
    将字符串切割成指定大小的块（带重叠），减少语义割裂。
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def read_docx(file_path):
    """
    读取 docx 文件，返回所有非空段落和表格单元格文本。
    """
    docx = DocxDocument(file_path)
    content = []
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
    return content

def read_pdf_with_pymupdf(file_path):
    """
    使用 PyMuPDF 读取 pdf 文本，并做基本预处理。
    """
    documents = []
    pdf = fitz.open(file_path)
    for i, page in enumerate(pdf):
        text = page.get_text("text")
        if text:
            processed_text = preprocess_text(text)
            documents.append(Document(text=processed_text, extra_info={"source": f"Page {i + 1}"}))
    pdf.close()
    return documents

def process_file(file_path, file_type, config):
    """
    根据文件类型读取文本，并使用 config 的 chunk_size, overlap 分块。
    返回分块后的 Document 列表。
    """
    if file_type == "docx":
        docx_content = read_docx(file_path)
        if not docx_content:
            raise ValueError("未从 docx 文件中读取到任何内容。")
        documents = [
            Document(text=preprocess_text(text), extra_info={"source": f"Content {i + 1}"})
            for i, text in enumerate(docx_content)
        ]
    elif file_type == "pdf":
        documents = read_pdf_with_pymupdf(file_path)
    elif file_type == "txt":
        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    # 分块
    chunked_documents = []
    for doc in documents:
        text_chunks = chunk_text_with_overlap(doc.text, config.chunk_size, config.overlap)
        for chunk in text_chunks:
            chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))
    return chunked_documents

def setup_storage_dir(file_name):
    """
    创建统一存储目录
    """
    base_dir = os.path.join(os.getcwd(), f"{file_name}_knowledge_graph")
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
    return base_dir

def check_subfolder_exists(parent_folder, subfolder_name):
    """
    检查指定文件夹中是否存在某个子文件夹
    """
    subfolder_path = os.path.join(parent_folder, subfolder_name)
    return os.path.exists(subfolder_path) and os.path.isdir(subfolder_path)