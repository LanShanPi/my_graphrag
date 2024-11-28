from llama_index.core.prompts import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType
import os
from llama_index.core import SimpleDirectoryReader, KnowledgeGraphIndex, StorageContext, load_index_from_storage
from llama_index.core import PropertyGraphIndex
from llama_index.core.graph_stores import SimpleGraphStore
from llama_index.llms.openai import OpenAI
from pyvis.network import Network
import IPython
import openai
import torch
from llama_index.readers.file import PyMuPDFReader
import pdfplumber
from llama_index.core.schema import Document
import re
from docx import Document as DocxDocument
from concurrent.futures import ThreadPoolExecutor
import pdfplumber
import fitz  # PyMuPDF
from config import OPENAI_API_KEY1,API_BASE


# 设置 OpenAI API
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY1
openai.api_key = os.environ["OPENAI_API_KEY"]
openai.api_base = API_BASE

# 配置 OpenAI 模型，使用 GPU 加速
# temperature: 控制生成文本的随机性，值越低，生成的文本越确定
# model: 指定使用的 OpenAI 模型
# device: 指定模型运行的设备（CPU 或 GPU）
# torch.cuda.is_available()目前返回false，所以用的是cpu
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")  # 检查是否有可用的 GPU
llm = OpenAI(temperature=0, model="gpt-4o", device=device)  # 初始化 OpenAI 模型



# 读取 .docx 文件内容，包括段落和表格
def read_docx(file_path):
    """
    读取 .docx 文件，包括普通段落和表格内容
    :param file_path: .docx 文件路径
    :return: 提取的内容列表
    """
    docx = DocxDocument(file_path)
    content = []

    # 读取普通段落
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)

    # 读取表格内容
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))  # 将一行的单元格合并成一条记录

    return content

# 文本预处理和分块函数
def preprocess_text(text):
    """
    对文本进行清理，比如去除多余空格或特殊字符
    :param text: 原始文本
    :return: 清理后的文本
    """
    text = re.sub(r'\s+', ' ', text).strip()  # 清除多余空格
    return text

def chunk_text_with_overlap(text, chunk_size, overlap):
    """
    分块文本，支持重叠
    :param text: 原始文本
    :param chunk_size: 每块字符数
    :param overlap: 重叠字符数
    :return: 分块后的文本列表
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def hongfushi():
    # 定义自定义提示模板
    CUSTOM_KG_TRIPLET_EXTRACT_TMPL = (
        "从给定文本中提取最多 {max_knowledge_triplets} 个知识三元组。"
        "每个三元组应以 (head, relation, tail) 及其各自的类型形式出现。\n"
        "---------------------\n"
        "初始本体论：\n"
        "实体类型：{allowed_entity_types}\n"
        "关系类型：{allowed_relation_types}\n"
        "\n"
        "以这些类型为起点，但根据上下文需要引入新类型。\n"
        "\n"
        "指南：\n"
        "- 以 JSON 格式输出：[{{'head': '', 'head_type': '', 'relation': '', 'tail': '', 'tail_type': ''}}]\n"
        "- 使用最完整的实体形式（例如，使用 'Red Fuji Apple' 而不是 'Apple'）\n"
        "- 保持实体简洁（最多 3-5 个词）\n"
        "- 将复杂短语分解为多个三元组\n"
        "- 确保知识图谱连贯且易于理解\n"
        "- 专注于红富士苹果生产的关键技术、病虫害防治、防灾减灾技术等方面\n"
        "- 包括品种特性、栽培方法、病虫害名称及其影响、防灾策略等\n"
        "---------------------\n"
        "文本：{text}\n"
        "输出：\n"
    )
    CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = PromptTemplate(
        CUSTOM_KG_TRIPLET_EXTRACT_TMPL, 
        prompt_type=PromptType.KNOWLEDGE_TRIPLET_EXTRACT
    )

    # 定义允许的实体类型和关系类型
    allowed_entity_types = [
        "PERSON", "COMPANY", "PRODUCT", "UNIVERSITY", "HEALTH_METRIC", 
        "RESEARCH_TOPIC", "FRUIT_VARIETY", "PESTICIDE", "DISEASE", "CULTIVATION_TECHNIQUE",
        "CLIMATIC_CONDITION", "SOIL_TYPE", "WATER_MANAGEMENT", "NUTRIENT_MANAGEMENT",
        "PEST_OR_DISEASE", "AGRICULTURAL_PRACTICE"
    ]
    allowed_relation_types = [
        "CEO_OF", "PRODUCES", "MONITORS", "STUDIES", "TREATS", "PREVENTS", 
        "USES", "INCLUDES", "MANAGES", "CAUSES", "AFFECTS", "REQUIRES",
        "IMPROVES", "REDUCES", "OPTIMIZES", "FOLLOWS"
    ]
    return allowed_entity_types,allowed_relation_types,CUSTOM_KG_TRIPLET_EXTRACT_PROMPT

def mingchao():
    CUSTOM_KG_TRIPLET_EXTRACT_TMPL = (
        "从给定文本中提取与明朝历史相关的知识三元组。"
        "每个三元组应以 (head, relation, tail) 及其各自的类型形式出现。\n"
        "---------------------\n"
        "初始本体论：\n"
        "实体类型：{allowed_entity_types}\n"
        "关系类型：{allowed_relation_types}\n"
        "\n"
        "以这些类型为起点，但根据上下文需要引入新类型。\n"
        "\n"
        "指南：\n"
        "- 以 JSON 格式输出：[{{'head': '', 'head_type': '', 'relation': '', 'tail': '', 'tail_type': ''}}]\n"
        "- 使用最完整的实体形式（例如，使用 '朱元璋' 而不是 '朱'）\n"
        "- 保持实体简洁（最多 3-5 个词）\n"
        "- 将复杂短语分解为多个三元组\n"
        "- 确保知识图谱连贯且易于理解\n"
        "- 专注于明朝的关键人物、事件、政策、战役、制度等方面\n"
        "- 包括人物的出生、登基、去世、战役的起因、过程、结果、政策的制定与影响等\n"
        "---------------------\n"
        "文本：{text}\n"
        "输出：\n"
        )

    CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = PromptTemplate(
        CUSTOM_KG_TRIPLET_EXTRACT_TMPL,
        prompt_type=PromptType.KNOWLEDGE_TRIPLET_EXTRACT
        )
    allowed_entity_types = [
    "PERSON", "EVENT", "POLITICAL_FIGURE", "MILITARY_LEADER",
    "DYNASTY", "BATTLE", "CAPITAL", "POLICY", "INSTITUTION",
    "CULTURAL_FIGURE", "ECONOMIC_POLICY", "SOCIAL_SYSTEM"
    ]

    allowed_relation_types = [
    "FOUNDER", "MEMBER", "GENERAL", "CAPTURED_BY", "SUCCEEDED_BY",
    "PARTICIPANT", "RESULTED_IN", "IMPLEMENTED", "INFLUENCED_BY",
    "CULTURAL_CONTRIBUTION", "ECONOMIC_IMPACT", "SOCIAL_REFORM",
    "BORN_IN", "DIED_IN", "ASCENDED_THRONE", "REMOVED_FROM_POWER",
    "STARTED_REIGN", "ENDED_REIGN", "LED_BATTLE", "DEFEATED",
    "ESTABLISHED_POLICY", "CREATED_INSTITUTION", "WROTE_WORKS",
    "ADVOCATED_POLICY", "INITIATED_REFORM"]

    return allowed_entity_types,allowed_relation_types,CUSTOM_KG_TRIPLET_EXTRACT_PROMPT



def read_pdf_with_pymupdf(file_path):
    """
    使用 PyMuPDF 高效读取 PDF 文本
    """
    documents = []
    pdf = fitz.open(file_path)
    for i, page in enumerate(pdf):
        text = page.get_text("text")
        if text:  # 确保页面有内容
            processed_text = preprocess_text(text)
            documents.append(Document(text=processed_text, extra_info={"source": f"Page {i + 1}"}))
    pdf.close()
    return documents

def process_file(file_path, file_type):
    if file_type == "docx":
        docx_content = read_docx(file_path)
        if not docx_content:
            raise ValueError("未从文件中读取到任何内容，请检查文件格式或内容是否为空。")
        documents = [
            Document(text=preprocess_text(text), extra_info={"source": f"Content {i + 1}"})
            for i, text in enumerate(docx_content)
        ]
    elif file_type == "pdf":
        documents = read_pdf_with_pymupdf(file_path)  # 或 read_pdf_parallel(file_path)
    elif file_type == "txt":
        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return documents

# 定义存储目录
# PERSIST_DIR = "honglou_knowledge_graph"
# HTML_FILE = "honglou_knowledge_graph.html"

# PERSIST_DIR = "pingguo_knowledge_graph"
# HTML_FILE = "pingguo_knowledge_graph.html"

PERSIST_DIR = "mingchao_knowledge_graph"
HTML_FILE = "mingchao_knowledge_graph.html"

# 检查存储目录是否存在
if not os.path.exists(PERSIST_DIR):
    # file_path = "/home/share/shucshqyfzyxgsi/home/lishuguang/my_graphrag/data/红富士苹果生产.docx"
    file_path = "/home/share/shucshqyfzyxgsi/home/lishuguang/my_graphrag/data/明朝那些事儿.pdf"
    documents = process_file(file_path,"pdf")
    # 分块文本
    chunk_size = 512  # 设置每个块的字符数
    overlap = 50  # 设置重叠长度
    chunked_documents = []
    for doc in documents:
        text_chunks = chunk_text_with_overlap(doc.text, chunk_size,overlap)
        for chunk in text_chunks:
            chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))

    # 创建存储上下文
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    
    # 创建 KnowledgeGraphIndex
    # max_triplets_per_chunk表示文档分块后每个块可以提取多少三元组（关系和事实）的数量
    # kg_triplet_extract_fn 允许用户根据特定的需求和数据特性调整三元组提取的方式，而不是依赖于默认的提取方法。
    # include_embeddings=True：语义搜索：通过包含嵌入，索引可以支持基于语义的查询，而不仅仅是关键词匹配。这使得查询能够返回语义上相关的结果，即使它们不包含查询中的确切词语。
                        # 相似性计算：嵌入允许计算不同文本之间的相似性，从而支持更复杂的查询和分析。
                        # 提高查询准确性：在某些应用场景中，嵌入可以提高查询的准确性和相关性，因为它们捕捉了文本的语义信息。
    # show_progress=True显示进度条

    #     # 创建 KnowledgeGraphIndex
    # kg_index = PropertyGraphIndex(
    #     kg_triple_extract_template=CUSTOM_KG_TRIPLET_EXTRACT_PROMPT,
    #     allowed_entity_types=allowed_entity_types,
    #     allowed_relation_types=allowed_relation_types
    # )
    entity_types,relation_types,CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = mingchao()
    index = PropertyGraphIndex.from_documents(
        chunked_documents,
        max_triplets_per_chunk=10,
        storage_context=storage_context,
        show_progress=True,
        include_embeddings=False,
        # kg_triplet_extract_fn=custom_extract_triplets,
        kg_triple_extract_template=CUSTOM_KG_TRIPLET_EXTRACT_PROMPT,
        allowed_entity_types=entity_types,
        allowed_relation_types=relation_types
    )
    # 存储索引
    storage_context.persist(persist_dir=PERSIST_DIR)
    if not os.path.exists(HTML_FILE):
        index.property_graph_store.save_networkx_graph(name=HTML_FILE)
else:
    # 加载已存储的索引
    storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
    index = load_index_from_storage(storage_context)
# 查询索引
# include_text 用途：决定在查询结果中是否包含原始文本。
# True：当设置为 True 时，查询结果将包含与匹配的三元组相关的完整文本块。这对于需要查看上下文或详细信息的应用场景非常有用。
# False：当设置为 False 时，查询结果只包含匹配的三元组，而不包括完整的文本块。这可以减少返回的数据量，提高查询速度。

query_engine = index.as_query_engine(llm=llm, include_text=False)
# query_engine.set_query_mode("compact")
response = query_engine.query("朱元璋当皇帝前都做过什么")
print(response)


