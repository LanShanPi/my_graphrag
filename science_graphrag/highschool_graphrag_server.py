import os
from config import OPENAI_API_KEY1 as OPENAI_API_KEY
from config import API_BASE3 as API_BASE
import openai


os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
openai.api_key = os.environ["OPENAI_API_KEY"]
openai.api_base = API_BASE


from llama_index.core.prompts import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import SimpleDirectoryReader, StorageContext, load_index_from_storage
from llama_index.core import PropertyGraphIndex
from llama_index.core.graph_stores import SimpleGraphStore
from llama_index.llms.openai import OpenAI as LlamaOpenAI
from pyvis.network import Network
import openai
from llama_index.core.schema import Document
import re
from docx import Document as DocxDocument
import fitz  # PyMuPDF
from prompt.prompt import science_general_template as triplet_extraction_template
from functools import lru_cache
import networkx as nx
from openai import OpenAI as local_openai
from prompt.response_prompt import mingchao_person as prompt_
import time
import logging
import hanlp
import fitz  # PyMuPDF
from paddleocr import PaddleOCR
from PIL import Image
import io
from sympy import sympify
from sympy.parsing.sympy_parser import parse_expr
import torch

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'  # 只显示错误，忽略警告

# 定义统一的存储目录
def setup_storage_dir(file_name):
    """
    创建统一存储目录
    :param file_name: 文件名称
    :return: 存储目录路径
    """
    base_dir = os.path.join(os.getcwd(), f"{file_name}_knowledge_graph")
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
    return base_dir

# 构建层次化图谱
def build_hierarchical_graph(triples, context):
    """
    根据学科、章节和知识点构建分层知识图谱
    """
    G = nx.DiGraph()
    for triple in triples:
        subject = context.get("subject", "General Science")
        section = triple["context"].get("section", "Unknown Section")
        
        # 添加学科节点
        if not G.has_node(subject):
            G.add_node(subject, label="Subject")
        
        # 添加章节节点
        if not G.has_node(section):
            G.add_node(section, label="Section")
            G.add_edge(subject, section, relation="CONTAINS")
        
        # 添加知识点节点
        head = triple["head"]
        tail = triple["tail"]
        relation = triple["relation"]
        
        if not G.has_node(head):
            G.add_node(head, label="Knowledge Point")
        if not G.has_node(tail):
            G.add_node(tail, label="Knowledge Point")
        
        # 添加知识点的关系边
        G.add_edge(head, tail, label=relation)
        G.add_edge(section, head, relation="CONTAINS")
    return G

# 文件读取函数
def read_docx(file_path):
    docx = DocxDocument(file_path)
    content = []
    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            content.append(text)
    for table in docx.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
    return content

def preprocess_text(text):
    # 将字符串 text 中的所有连续空白字符替换为单个空格。
    return re.sub(r'\s+', ' ', text).strip()

# def chunk_text_with_overlap(text, chunk_size, overlap):
#     # 将字符串切割成指定大小的块
#     chunks = []
#     start = 0
#     while start < len(text):
#         end = start + chunk_size
#         chunks.append(text[start:end])
#         start += chunk_size - overlap
#     return chunks


# def protect_formulas(text):
#     """
#     使用正则表达式识别并保护公式
#     """
#     formula_pattern = r"[A-Za-z0-9+\-*/^=(){}[\].]+"
#     formulas = re.findall(formula_pattern, text)
#     formula_map = {f"{{FORMULA_{i}}}": formula for i, formula in enumerate(formulas)}

#     # 替换公式为占位符
#     for placeholder, formula in formula_map.items():
#         text = text.replace(formula, placeholder)
#     return text, formula_map

def protect_formulas_advanced(text):
    """
    高级公式保护：结合改进正则表达式和 SymPy 解析
    """
    formula_pattern = r"""
    (?:                                     # 非捕获组，匹配以下之一
        [a-zA-Z0-9_+\-*/^=<>≤≥(),\[\]{}|]+  # 基本符号、字母、数字
        | [√∫∞ΣΠΔ∂⋅→←↔↕∝∑∏θλμνξορστυφχψω]  # 常见希腊字母及数学符号
        | \d+\.\d+(e[+\-]?\d+)?             # 科学计数法（如 1.23e-4）
        | \d+\.\d+                          # 带小数点的数字（如 3.14）
        | [A-Z][a-z]?\d*(?:[+\-]?)          # 化学元素（如 H2O、CO2）
        | [+\-]?\d+\s*[a-zA-Z]{1,3}         # 单位（如 5 m、10 kg）
        | [a-zA-Z]+\([a-zA-Z0-9, ]+\)       # 数学函数（如 sin(x) 或 log(x, y)）
        | [a-zA-Z0-9]+(?:_\d+|^\d+)?        # 下标或上标（如 x_1, x^2）
        | ->|<-|⇌|=                         # 化学反应符号
        | [ATCG]+                           # DNA 序列
    )
    """
    formulas = re.findall(formula_pattern, text)
    formula_map = {}

    for i, formula in enumerate(formulas):
        try:
            # 使用 SymPy 解析公式
            parsed_formula = sympify(formula)
            placeholder = f"{{FORMULA_{i}}}"
            formula_map[placeholder] = str(parsed_formula)
            text = text.replace(formula, placeholder)
        except Exception:
            pass  # 如果解析失败，跳过
    return text, formula_map

def restore_formulas(chunks, formula_map):
    """
    在分块后还原公式
    """
    restored_chunks = []
    for chunk in chunks:
        for placeholder, formula in formula_map.items():
            chunk = chunk.replace(placeholder, formula)
        restored_chunks.append(chunk)
    return restored_chunks

def semantic_chunking_with_context(text, chunk_size, overlap):
    """
    对理科文本进行语义分块，支持公式保护和上下文信息附加
    """
    sent_splitter = hanlp.load('CTB6_CONVSEG')

    # 章节标题正则匹配
    section_pattern = r"(第[一二三四五六七八九十百]+[章节]|[0-9]+\.[0-9]+.*?节)"
    matches = re.finditer(section_pattern, text)

    sections = []
    last_pos = 0
    current_section = "Unknown Section"

    # 提取章节信息
    for match in matches:
        if last_pos != 0:
            sections.append((current_section, text[last_pos:match.start()]))
        current_section = match.group(0)
        last_pos = match.start()
    sections.append((current_section, text[last_pos:]))

    chunks = []
    for section, section_text in sections:
        # 保护公式
        protected_text, formula_map = protect_formulas_advanced(section_text)
        # 分句
        sentences = sent_splitter(protected_text)
        # 分块
        current_chunk = []
        current_length = 0
        for sent in sentences:
            if current_length + len(sent) > chunk_size:
                # 保存当前块
                chunks.append({
                    "text": "".join(current_chunk),
                    "section": section,
                    "formula_map": formula_map  # 保留公式映射
                })
                # 重叠部分
                current_chunk = current_chunk[-(overlap // len(sent)):] if current_chunk else []
                current_length = len("".join(current_chunk))
            current_chunk.append(sent)
            current_length += len(sent)

        # 添加最后一个块
        if current_chunk:
            chunks.append({
                "text": "".join(current_chunk),
                "section": section,
                "formula_map": formula_map
            })

    # 还原公式
    restored_chunks = []
    for chunk in chunks:
        chunk["text"] = restore_formulas([chunk["text"]], chunk["formula_map"])[0]
        restored_chunks.append(chunk)

    return restored_chunks


# def read_pdf_with_pymupdf(file_path):
#     documents = []
#     pdf = fitz.open(file_path)
#     for i, page in enumerate(pdf):
#         text = page.get_text("text")
#         if text:
#             processed_text = preprocess_text(text)
#             documents.append(Document(text=processed_text, extra_info={"source": f"Page {i + 1}"}))
#     pdf.close()
#     return documents

def read_pdf_with_pymupdf(file_path):
    """
    优化后的 PDF 处理模块：
    1. 提取嵌入式文本。
    2. 对图片中的文字进行 OCR 识别。
    3. 返回处理后的 Document 列表。
    """
    documents = []
    pdf = fitz.open(file_path)  # 打开 PDF 文件

    for i, page in enumerate(pdf):
        # Step 1: 提取嵌入式文本
        text = page.get_text("text")
        text = preprocess_text(text)

        # 如果文本为空，尝试从页面图片中提取文字
        if not text.strip():
            text = extract_text_with_paddleocr(page)
        # 如果仍然没有文字，标记为无法识别
        if not text.strip():
            text = f"无法提取文字，来源：第 {i + 1} 页"
        # 创建 Document 对象
        documents.append(Document(text=text, extra_info={"source": f"Page {i + 1}"}))

    pdf.close()
    return documents

def extract_text_with_paddleocr(page):
    """
    使用 PaddleOCR 从 PDF 页面中的图片提取文字
    :param page: PDF 页面对象
    :return: 提取的文字
    """
    ocr = PaddleOCR(use_angle_cls=True, lang="ch")
    ocr_text = []
    image_list = page.get_images(full=True)

    for img_index, img in enumerate(image_list):
        xref = img[0]
        base_image = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False, xref=xref)  # 提高清晰度
        image = Image.open(io.BytesIO(base_image.tobytes("png")))  # 转换为 PIL Image 对象

        # 使用 PaddleOCR 识别文字
        ocr_results = ocr.ocr(image, cls=True)
        for line in ocr_results[0]:
            ocr_text.append(line[1][0])  # 提取文字内容

    return "\n".join(ocr_text)

def process_file(file_path, file_type):
    if file_type == "docx":
        docx_content = read_docx(file_path)
        if not docx_content:
            raise ValueError("未从文件中读取到任何内容，请检查文件格式或内容是否为空。")
        documents = [
            Document(text=preprocess_text(text), extra_info={"source": f"Content {i + 1}"})
            for i, text in enumerate(docx_content)
        ]
    elif file_type == "pdf":
        documents = read_pdf_with_pymupdf(file_path)
    elif file_type == "txt":
        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return documents

# 定义知识提取模板
def get_config():
    CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = PromptTemplate(
        triplet_extraction_template["CUSTOM_KG_TRIPLET_EXTRACT_TMPL"],
        prompt_type=PromptType.KNOWLEDGE_TRIPLET_EXTRACT
    )
    allowed_entity_types = triplet_extraction_template["allowed_entity_types"]
    allowed_relation_types = triplet_extraction_template["allowed_relation_types"]
    return allowed_entity_types, allowed_relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT

def check_subfolder_exists(parent_folder, subfolder_name):
    """
    检查指定文件夹中是否存在某个子文件夹
    :param parent_folder: 父文件夹路径
    :param subfolder_name: 子文件夹名称
    :return: True 如果存在，False 如果不存在
    """
    subfolder_path = os.path.join(parent_folder, subfolder_name)
    return os.path.exists(subfolder_path) and os.path.isdir(subfolder_path)

def generate_knowledge_graph(file_path, file_type, dir_name, storage_dir):
    # File processing and knowledge graph construction
    documents = process_file(file_path, file_type)
    chunk_size = 1024
    overlap = 256
    chunked_documents = []
    # for doc in documents:
    #     text_chunks = semantic_chunking_with_context(doc.text, chunk_size, overlap)
    #     for chunk in text_chunks:
    #         chunked_documents.append(Document(text=chunk, extra_info=doc.extra_info))


    for doc in documents:
        text_chunks = semantic_chunking_with_context(doc.text, chunk_size, overlap)
        for chunk in text_chunks:
            # 只传递 chunk["text"]，并将上下文信息附加到 extra_info
            chunked_documents.append(Document(
                text=chunk["text"],
                extra_info={**doc.extra_info, "section": chunk["section"]}
            ))

    # Load templates
    entity_types, relation_types, CUSTOM_KG_TRIPLET_EXTRACT_PROMPT = get_config()

    # Create storage context
    graph_store = SimpleGraphStore()
    storage_context = StorageContext.from_defaults(graph_store=graph_store)
    # text-embedding-ada-002:通用性强，适合自然语言与公式混合的文本;SciBERT:针对科学文本优化，能理解公式上下文
    embed_model = OpenAIEmbedding(model_name="text-embedding-ada-002",api_key=OPENAI_API_KEY, api_base=API_BASE)
    # embed_model=embed_model,  # 指定编码模型
    # include_embeddings=True,

    # Build knowledge graph
    index = PropertyGraphIndex.from_documents(
        chunked_documents,
        llm=llm,
        embed_model=embed_model,  # 指定编码模型
        include_embeddings=True,
        max_triplets_per_chunk=10,
        storage_context=storage_context,
        show_progress=True,
        kg_triple_extract_template=CUSTOM_KG_TRIPLET_EXTRACT_PROMPT,
        allowed_entity_types=entity_types,
        allowed_relation_types=relation_types,
    )

    # Store knowledge graph and visualization files
    index_dir = os.path.join(storage_dir, "index")
    if not os.path.exists(index_dir):
        os.makedirs(index_dir)
    storage_context.persist(persist_dir=index_dir)

    # Convert knowledge graph to NetworkX graph
    G = nx.DiGraph()
    for triplet in index.property_graph_store.get_triplets():
        subj = triplet["head"]
        obj = triplet["tail"]
        rel = triplet["relation"]
        time = triplet.get("time", "未知时间")  # Get time or set default
        
        # Add nodes and edges to the graph
        G.add_node(subj, label=subj)
        G.add_node(obj, label=obj)
        # Add time to the edge label
        G.add_edge(subj, obj, label=f"{rel} ({time})")

    # Visualize with pyvis
    net = Network(notebook=True, cdn_resources="in_line", directed=True)
    net.from_nx(G)

    # Save HTML file
    html_file = os.path.join(storage_dir, f"{dir_name}_graph.html")
    net.show(html_file)
    print(f"HTML file has been generated: {html_file}")
    return index

def load_knowledge_graph(storage_dir):
    # 加载已存储的索引
    print("已存在相关知识库")
    storage_context = StorageContext.from_defaults(persist_dir=storage_dir+"/"+"index")
    index = load_index_from_storage(storage_context)
    return index


def get_response(index,queries):
    ### 普通回复
    # 查询索引
    # include_text 用途：决定在查询结果中是否包含原始文本。
    # True：当设置为 True 时，查询结果将包含与匹配的三元组相关的完整文本块。这对于需要查看上下文或详细信息的应用场景非常有用。
    # False：当设置为 False 时，查询结果只包含匹配的三元组，而不包括完整的文本块。这可以减少返回的数据量，提高查询速度。
    ######
    # response_mode 用途：指定如何生成和组织查询响应。
    # tree_summarize：在这种模式下，系统会递归地构建一个树状结构来总结查询结果。这种模式适用于需要对大量信息进行总结的场景。
    # 工作原理：tree_summarize 会将所有相关的文本块合并，并通过多次调用 LLM 来生成一个总结，最终返回一个简洁的响应。
    # 适用场景：适合需要高层次总结的场景，结果更全面，但是速度可能慢

    # refine 模式用途：用于逐步改进和细化响应。
    # 工作原理：在 refine 模式下，系统会首先使用第一个文本块和查询生成一个初始答案。然后，它会将这个答案与下一个文本块结合，再次生成一个改进的答案。这个过程会持续进行，直到所有文本块都被处理完毕。
    # 适用场景：适合需要详细和精确答案的场景，因为它会逐步整合所有相关信息。

    # compact 模式用途：用于在生成响应之前合并文本块，以减少对 LLM 的调用次数。
    # 工作原理：在 compact 模式下，系统会尽可能多地将文本块合并到一个上下文窗口中，然后生成响应。这种方法减少了对 LLM 的调用次数，因为它在合并后的文本上进行处理。
    # 适用场景：适合需要快速响应的场景，因为它通过减少 LLM 调用次数来提高效率。

    # 这两种模式各有优劣，refine 提供更详细的答案，而 compact 提供更快的响应速度。

    # vector_store_query_mode：指定向量存储查询的模式，通常用于选择不同的向量检索策略。
        # DEFAULT：默认的向量搜索模式，通常用于基于向量相似度的检索。
        # SPARSE：稀疏向量搜索模式，适用于稀疏数据的检索。
        # HYBRID：混合搜索模式，结合了向量相似度和文本搜索。
        # TEXT_SEARCH：全文本搜索模式，基于文本内容进行检索。
        # SEMANTIC_HYBRID：语义混合模式，结合语义理解和向量相似度。

    # filters：用于在查询时应用元数据过滤器。可以根据文档的元数据（如作者、日期等）来筛选符合条件的节点。

    # # 定义元数据过滤器（只检索作者为“历史学家”在2023年的记录，key的值要根据原数据中的标记来设置，查看原数据中的标签暂未写）####
    # filters = MetadataFilters(
    #                 filters=[
    #                     ExactMatchFilter(key="author", value="历史学家"),
    #                     ExactMatchFilter(key="year", value=2023)
    #                     ]
    #                     )

    query_engine = index.as_query_engine(llm=llm, include_text=True,response_mode="tree_summarize",similarity_top_k=20,vector_store_query_mode="SEMANTIC_HYBRID")
    response = query_engine.query(queries)
    for idx, source in enumerate(response.source_nodes):
        # 使用 get_content() 方法获取节点内容
        print("[Source] " + str(idx) + ": ", source.node.get_content())

    return response


if __name__ == "__main__":
    # 配置 GPU/CPU 设备
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    llm = LlamaOpenAI(temperature=0, model="gpt-4o", api_key=OPENAI_API_KEY, api_base=API_BASE,timeout=600)

    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler("knowledge_graph.log"),
            logging.StreamHandler()
        ]
    )


    file_path = "/data/hongzhili/my_graphrag/data/高中物理测试题3.pdf"
    file_type = "pdf"
    dir_name = "physics"
    
    # 生成存储路径
    storage_dir = setup_storage_dir(dir_name)
    # 若未生成相关图谱，则进行生成，若之前已生成则直接加载
    if not check_subfolder_exists(storage_dir, "index"):
        index = generate_knowledge_graph(file_path,file_type,dir_name,storage_dir)
    else:
        index = load_knowledge_graph(storage_dir)

    input_ = None
    # 不要生成 LaTeX 语法。这个提示很重要
    pre_prompt = "根据自身能力和检索到的知识尽可能详细的回复下述问题，且回复要满足：回答的准确性、回答的完整性、回答的逻辑性、回答的语言表达清晰性，这四个要求，仅需输出回答就好了，不需要额外的输出。不要生成 LaTeX 语法。下面是问题："
    while input_ != "over":
        input_ = input("请输入你针对’高中物理测试题3的‘这本书的问题：")
        start = time.time()
        print("#######################")
        print("问题：",input_)
        response = get_response(index,pre_prompt+input_)
        print("#######################")
        print("生成时间：",time.time()-start)
        print(response)
